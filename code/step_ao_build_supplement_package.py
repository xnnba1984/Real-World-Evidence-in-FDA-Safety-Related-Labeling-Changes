#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
BASE_DIR = REPO_ROOT
ANALYSIS_OUTPUTS = BASE_DIR / "analysis_outputs"
MANUSCRIPT_DIR = BASE_DIR / "manuscript"
SUPPLEMENT_OUTPUT_DIR = ANALYSIS_OUTPUTS / "supplement_package"
SUPPLEMENT_FIGURE_DIR = ANALYSIS_OUTPUTS / "figures" / "supplement"

TABLE_PATHS = {
    "table_s1": ANALYSIS_OUTPUTS / "supplement_year_application_type_outcomes.csv",
    "table_s2": ANALYSIS_OUTPUTS / "figure4_explicitness_counts.csv",
    "table_s3": ANALYSIS_OUTPUTS / "sensitivity_prevalence_tables.csv",
    "table_s5": ANALYSIS_OUTPUTS / "product_age" / "product_age_outcomes.csv",
    "table_s6": ANALYSIS_OUTPUTS / "therapeutic_area" / "therapeutic_area_outcomes.csv",
    "table_s7": ANALYSIS_OUTPUTS / "sponsor_manufacturer" / "sponsor_manufacturer_outcomes.csv",
}
VALIDATION_SAMPLE_PATH = BASE_DIR / "analysis_ready" / "human_validation_package" / "validation_sample_master.csv"

SUPPLEMENT_APPENDIX_MD = MANUSCRIPT_DIR / "supplementary_appendix_v1.md"
SUPPLEMENT_APPENDIX_DOCX = MANUSCRIPT_DIR / "supplementary_appendix_v1.docx"
SUPPLEMENT_INVENTORY_MD = MANUSCRIPT_DIR / "supplement_inventory.md"
SUPPLEMENT_INVENTORY_CSV = MANUSCRIPT_DIR / "supplement_inventory.csv"

TABLE_CELL_FONT_SIZE = Pt(11)
BODY_FONT_SIZE = Pt(12)
CAPTION_FONT_SIZE = Pt(11)
LANDSCAPE_MARGINS = Inches(0.6)
FIGURE_WIDTH = Inches(9.3)
TABLE_TOTAL_WIDTH = 9.4


def round_series(series: pd.Series, decimals: int = 1) -> pd.Series:
    return series.astype(float).round(decimals)


def clean_annual_trends(df: pd.DataFrame) -> pd.DataFrame:
    outcome_order = {
        "Broad public RWE": 0,
        "Analytic public RWE": 1,
        "Explicit public RWE": 2,
    }
    cleaned = df.rename(
        columns={
            "event_year": "Event year",
            "application_type": "Application type",
            "endpoint_label": "Outcome",
            "count_yes": "Positive events",
            "denominator": "Events in stratum",
            "pct_yes": "Prevalence (%)",
        }
    ).copy()
    cleaned["Outcome"] = cleaned["Outcome"].replace({"Main public RWE": "Broad public RWE"})
    cleaned["Prevalence (%)"] = round_series(cleaned["Prevalence (%)"])
    cleaned["_outcome_order"] = cleaned["Outcome"].map(outcome_order).fillna(99)
    cleaned = cleaned.sort_values(["Event year", "Application type", "_outcome_order"]).drop(columns=["_outcome_order"])
    return cleaned[
        [
            "Event year",
            "Application type",
            "Outcome",
            "Positive events",
            "Events in stratum",
            "Prevalence (%)",
        ]
    ].reset_index(drop=True)


def clean_explicitness(df: pd.DataFrame) -> pd.DataFrame:
    tier_map = {
        "no_public_basis_found": "No public basis found",
        "unclear_public_basis": "Unclear public basis",
        "spontaneous_reports_only": "Spontaneous reports only",
        "explicit_observational_real_world": "Explicit observational real-world evidence",
        "explicit_rwe": "Explicit real-world evidence",
    }
    subset_map = {"overall": "Overall", "BLA": "BLA", "NDA": "NDA"}
    subset_order = {"Overall": 0, "BLA": 1, "NDA": 2}
    tier_order = {
        "No public basis found": 0,
        "Unclear public basis": 1,
        "Spontaneous reports only": 2,
        "Explicit observational real-world evidence": 3,
        "Explicit real-world evidence": 4,
    }
    cleaned = df.copy()
    cleaned["Subset"] = cleaned["group"].map(subset_map).fillna(cleaned["group"])
    cleaned["Evidence explicitness"] = cleaned["evidence_explicitness_tier"].map(tier_map).fillna(
        cleaned["evidence_explicitness_tier"]
    )
    cleaned["Percent (%)"] = round_series(cleaned["pct"])
    cleaned = cleaned.rename(columns={"count": "Count", "denominator": "Denominator"})
    cleaned["_subset_order"] = cleaned["Subset"].map(subset_order).fillna(99)
    cleaned["_tier_order"] = cleaned["Evidence explicitness"].map(tier_order).fillna(99)
    cleaned = cleaned.sort_values(["_subset_order", "_tier_order"]).drop(columns=["_subset_order", "_tier_order"])
    return cleaned[["Subset", "Evidence explicitness", "Count", "Denominator", "Percent (%)"]].reset_index(drop=True)


def clean_sensitivity(df: pd.DataFrame) -> pd.DataFrame:
    scenario_group_map = {
        "baseline_endpoint_ladder": "Measure definitions",
        "provenance_stratification": "Annotation provenance",
        "subset_restriction": "Subset restrictions",
    }
    scenario_name_map = {
        "baseline_full_cohort": "Baseline full cohort",
        "provenance_adjudication": "Adjudication rows",
        "provenance_first_pass": "First-pass rows",
        "provenance_repair": "Repair rows",
        "subset_confidence_high_medium": "High- or medium-confidence rows",
        "subset_confidence_high_only": "High-confidence rows",
        "subset_duplicate_collapsed": "Duplicate-cluster collapsed",
        "subset_extracted_docs_only": "Extracted-document rows",
        "subset_warning_excluded": "Aggressive warning exclusion",
        "subset_hard_issue_excluded": "QC-strict exclusion of hard-issue rows",
    }
    outcome_order = {
        "Broad public RWE measure": 0,
        "Analytic public RWE measure": 1,
        "Explicit public RWE sensitivity measure": 2,
        "Public RWE excluding spontaneous-report-only evidence": 3,
        "QC-strict broad public RWE sensitivity measure": 4,
    }
    cleaned = df.copy()
    cleaned["Scenario group"] = cleaned["scenario_group"].map(scenario_group_map).fillna(cleaned["scenario_group"])
    cleaned["Scenario"] = cleaned["scenario_name"].map(scenario_name_map).fillna(cleaned["scenario_label"])
    cleaned["Outcome"] = cleaned["outcome_label"]
    cleaned["Outcome"] = cleaned["Outcome"].replace(
        {
            "Main public RWE endpoint": "Broad public RWE measure",
            "Analytic public RWE endpoint": "Analytic public RWE measure",
            "Explicit public RWE sensitivity endpoint": "Explicit public RWE sensitivity measure",
            "Non-spontaneous public RWE sensitivity endpoint": "Public RWE excluding spontaneous-report-only evidence",
            "QC-strict public RWE sensitivity endpoint": "QC-strict broad public RWE sensitivity measure",
        }
    )
    cleaned = cleaned.rename(
        columns={
            "n_obs": "Events",
            "n_yes": "Positive events",
            "pct_yes": "Prevalence (%)",
            "delta_pct_points_from_baseline": "Change vs baseline (pp)",
            "relative_ratio_vs_baseline": "Ratio vs baseline",
        }
    )
    cleaned["Prevalence (%)"] = round_series(cleaned["Prevalence (%)"])
    cleaned["Change vs baseline (pp)"] = cleaned["Change vs baseline (pp)"].astype(float).round(1)
    cleaned["Ratio vs baseline"] = cleaned["Ratio vs baseline"].astype(float).round(2)
    cleaned["_group_order"] = cleaned["Scenario group"].map(
        {"Measure definitions": 0, "Subset restrictions": 1, "Annotation provenance": 2}
    ).fillna(99)
    cleaned["_scenario_order"] = cleaned["Scenario"].map(
        {
            "Baseline full cohort": 0,
            "QC-strict exclusion of hard-issue rows": 1,
            "Extracted-document rows": 2,
            "High- or medium-confidence rows": 3,
            "High-confidence rows": 4,
            "Duplicate-cluster collapsed": 5,
            "Aggressive warning exclusion": 6,
            "Adjudication rows": 7,
            "First-pass rows": 8,
            "Repair rows": 9,
        }
    ).fillna(99)
    cleaned["_outcome_order"] = cleaned["Outcome"].map(outcome_order).fillna(99)
    cleaned = cleaned.sort_values(["_group_order", "_scenario_order", "_outcome_order"])
    cleaned = cleaned.drop(columns=["_group_order", "_scenario_order", "_outcome_order"])
    return cleaned[
        [
            "Scenario group",
            "Scenario",
            "Outcome",
            "Events",
            "Positive events",
            "Prevalence (%)",
            "Change vs baseline (pp)",
            "Ratio vs baseline",
        ]
    ].reset_index(drop=True)


def build_validation_sample_characteristics(df: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    total_n = len(df)
    domain_order = [
        ("Sampling stratum", "validation_primary_stratum", {
            "negative_reference": "Negative reference",
            "analytic_positive": "Analytic positive",
            "primary_positive_nonanalytic": "Primary positive, nonanalytic",
            "complex_case": "Complex case",
        }),
        ("Annotation confidence", "annotation_confidence", {
            "high": "High",
            "medium": "Medium",
            "low": "Low",
        }),
        ("Final label source", "final_label_source", {
            "first_pass": "First pass",
            "repair": "Repair",
            "adjudication": "Adjudication",
        }),
        ("Application type", "Application Type", {"NDA": "NDA", "BLA": "BLA"}),
        ("Hard issue flag", "hard_issue_flag", {"no": "No", "yes": "Yes"}),
    ]
    value_order = {
        "Sampling stratum": ["Negative reference", "Analytic positive", "Primary positive, nonanalytic", "Complex case"],
        "Annotation confidence": ["High", "Medium", "Low"],
        "Final label source": ["First pass", "Repair", "Adjudication"],
        "Application type": ["NDA", "BLA"],
        "Hard issue flag": ["No", "Yes"],
    }
    for domain_label, column_name, value_map in domain_order:
        counts = df[column_name].value_counts(dropna=False)
        for raw_value, count in counts.items():
            value = "Missing" if pd.isna(raw_value) else value_map.get(raw_value, str(raw_value))
            rows.append(
                {
                    "Domain": domain_label,
                    "Value": value,
                    "Count": int(count),
                    "Percent of sample (%)": round(100.0 * float(count) / float(total_n), 1),
                }
            )
    result = pd.DataFrame(rows)
    result["_domain_order"] = result["Domain"].map({name: idx for idx, (name, _, _) in enumerate(domain_order)})
    result["_value_order"] = result.apply(
        lambda row: value_order.get(row["Domain"], []).index(row["Value"])
        if row["Value"] in value_order.get(row["Domain"], [])
        else 99,
        axis=1,
    )
    result = result.sort_values(["_domain_order", "_value_order"]).drop(columns=["_domain_order", "_value_order"])
    return result.reset_index(drop=True)


def clean_product_age(df: pd.DataFrame) -> pd.DataFrame:
    grouping_map = {"product_age_band": "Product age band", "approval_era": "Approval era"}
    group_order = {
        "Product age band": ["0-4 years", "5-9 years", "10-19 years", "20-29 years", "30+ years", "Unknown"],
        "Approval era": ["Pre-1990", "1990s", "2000s", "2010s", "2020s", "Unknown"],
    }
    cleaned = df.rename(
        columns={
            "group_label": "Group",
            "n_events": "Events",
            "pct_all_events": "Percent of all events (%)",
            "main_public_rwe_yes": "Broad public RWE positive",
            "main_public_rwe_pct": "Broad public RWE prevalence (%)",
            "analytic_public_rwe_yes": "Analytic public RWE positive",
            "analytic_public_rwe_pct": "Analytic public RWE prevalence (%)",
            "mean_transparency_score": "Mean transparency score",
        }
    ).copy()
    cleaned["Grouping"] = cleaned["group_type"].map(grouping_map).fillna(cleaned["group_type"])
    for column in [
        "Percent of all events (%)",
        "Broad public RWE prevalence (%)",
        "Analytic public RWE prevalence (%)",
        "Mean transparency score",
    ]:
        cleaned[column] = round_series(cleaned[column])
    cleaned["_grouping_order"] = cleaned["Grouping"].map({"Product age band": 0, "Approval era": 1}).fillna(99)
    cleaned["_group_order"] = cleaned.apply(
        lambda row: group_order.get(row["Grouping"], []).index(row["Group"])
        if row["Group"] in group_order.get(row["Grouping"], [])
        else 99,
        axis=1,
    )
    cleaned = cleaned.sort_values(["_grouping_order", "_group_order"]).drop(columns=["_grouping_order", "_group_order"])
    return cleaned[
        [
            "Grouping",
            "Group",
            "Events",
            "Percent of all events (%)",
            "Broad public RWE positive",
            "Broad public RWE prevalence (%)",
            "Analytic public RWE positive",
            "Analytic public RWE prevalence (%)",
            "Mean transparency score",
        ]
    ].reset_index(drop=True)


def clean_therapeutic_area(df: pd.DataFrame) -> pd.DataFrame:
    cleaned = df.rename(
        columns={
            "therapeutic_area_label": "Therapeutic area",
            "n_events": "Events",
            "pct_all_events": "Percent of all events (%)",
            "main_public_rwe_yes": "Broad public RWE positive",
            "main_public_rwe_pct": "Broad public RWE prevalence (%)",
            "analytic_public_rwe_yes": "Analytic public RWE positive",
            "analytic_public_rwe_pct": "Analytic public RWE prevalence (%)",
            "mean_transparency_score": "Mean transparency score",
        }
    ).copy()
    for column in [
        "Percent of all events (%)",
        "Broad public RWE prevalence (%)",
        "Analytic public RWE prevalence (%)",
        "Mean transparency score",
    ]:
        cleaned[column] = round_series(cleaned[column])
    return cleaned[
        [
            "Therapeutic area",
            "Events",
            "Percent of all events (%)",
            "Broad public RWE positive",
            "Broad public RWE prevalence (%)",
            "Analytic public RWE positive",
            "Analytic public RWE prevalence (%)",
            "Mean transparency score",
        ]
    ].reset_index(drop=True)


def clean_sponsor_manufacturer(df: pd.DataFrame) -> pd.DataFrame:
    grouping_map = {
        "sponsor_manufacturer_structure": "Sponsor/manufacturer structure",
        "generic_biosimilar_like_flag": "Generic or biosimilar-like company flag",
    }
    cleaned = df.rename(
        columns={
            "group_label": "Group",
            "n_events": "Events",
            "pct_all_events": "Percent of all events (%)",
            "main_public_rwe_yes": "Broad public RWE positive",
            "main_public_rwe_pct": "Broad public RWE prevalence (%)",
            "analytic_public_rwe_yes": "Analytic public RWE positive",
            "analytic_public_rwe_pct": "Analytic public RWE prevalence (%)",
            "mean_transparency_score": "Mean transparency score",
        }
    ).copy()
    cleaned["Grouping"] = cleaned["group_type"].map(grouping_map).fillna(cleaned["group_type"])
    for column in [
        "Percent of all events (%)",
        "Broad public RWE prevalence (%)",
        "Analytic public RWE prevalence (%)",
        "Mean transparency score",
    ]:
        cleaned[column] = round_series(cleaned[column])
    cleaned["_grouping_order"] = cleaned["Grouping"].map(
        {"Sponsor/manufacturer structure": 0, "Generic or biosimilar-like company flag": 1}
    ).fillna(99)
    cleaned = cleaned.sort_values(["_grouping_order", "Group"]).drop(columns=["_grouping_order"])
    return cleaned[
        [
            "Grouping",
            "Group",
            "Events",
            "Percent of all events (%)",
            "Broad public RWE positive",
            "Broad public RWE prevalence (%)",
            "Analytic public RWE positive",
            "Analytic public RWE prevalence (%)",
            "Mean transparency score",
        ]
    ].reset_index(drop=True)


def markdown_table(df: pd.DataFrame) -> str:
    header = "| " + " | ".join(str(col) for col in df.columns) + " |"
    divider = "| " + " | ".join("---" for _ in df.columns) + " |"
    rows = []
    for row in df.itertuples(index=False, name=None):
        rows.append("| " + " | ".join(str(value).replace("|", "\\|") for value in row) + " |")
    return "\n".join([header, divider, *rows])


def build_inventory(items: list[dict[str, str]]) -> pd.DataFrame:
    return pd.DataFrame(items)[["Item", "Type", "Title", "Purpose", "Source file"]]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = LANDSCAPE_MARGINS
    section.bottom_margin = LANDSCAPE_MARGINS
    section.left_margin = LANDSCAPE_MARGINS
    section.right_margin = LANDSCAPE_MARGINS

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = BODY_FONT_SIZE
    normal_paragraph = style.paragraph_format
    normal_paragraph.space_after = Pt(6)
    normal_paragraph.line_spacing = 1.15


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = BODY_FONT_SIZE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)


def add_caption(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.name = "Times New Roman"
    label_run.font.size = CAPTION_FONT_SIZE
    text_run = paragraph.add_run(" " + text)
    text_run.font.name = "Times New Roman"
    text_run.font.size = CAPTION_FONT_SIZE


def add_figure(document: Document, image_path: Path) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=FIGURE_WIDTH)


def estimate_column_widths(df: pd.DataFrame) -> list[float]:
    weights: list[float] = []
    sample = df.head(30)
    for column in df.columns:
        values = [str(column), *sample[column].astype(str).tolist()]
        max_len = max(len(value) for value in values)
        weights.append(float(min(max(max_len, 8), 28)))
    total_weight = sum(weights)
    return [TABLE_TOTAL_WIDTH * weight / total_weight for weight in weights]


def add_dataframe_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(table)
    widths = estimate_column_widths(df)

    header_cells = table.rows[0].cells
    for idx, column in enumerate(df.columns):
        header_cells[idx].text = str(column)
        set_cell_width(header_cells[idx], widths[idx])
        set_cell_shading(header_cells[idx], "D9E2F3")
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = TABLE_CELL_FONT_SIZE
            run.bold = True

    for row_values in df.itertuples(index=False, name=None):
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = str(value)
            set_cell_width(row_cells[idx], widths[idx])
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = row_cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = TABLE_CELL_FONT_SIZE

    document.add_paragraph()


def merge_context_cells(table, column_indexes: list[int]) -> None:
    for column_index in column_indexes:
        row_index = 1
        while row_index < len(table.rows):
            base_text = table.cell(row_index, column_index).text.strip()
            if not base_text:
                row_index += 1
                continue
            end_index = row_index
            while end_index + 1 < len(table.rows):
                next_text = table.cell(end_index + 1, column_index).text.strip()
                if next_text != base_text:
                    break
                end_index += 1
            if end_index > row_index:
                for clear_index in range(row_index + 1, end_index + 1):
                    table.cell(clear_index, column_index).text = ""
                merged = table.cell(row_index, column_index).merge(table.cell(end_index, column_index))
                while len(merged.paragraphs) > 1:
                    merged._tc.remove(merged.paragraphs[-1]._p)
                if not merged.paragraphs:
                    merged.add_paragraph()
                merged.paragraphs[0].text = base_text
                merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in merged.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = TABLE_CELL_FONT_SIZE
            row_index = end_index + 1


def build_docx_appendix(items: list[dict[str, object]]) -> None:
    document = Document()
    configure_document(document)
    add_heading(document, "Supplementary Materials", level=1)
    add_body_paragraph(
        document,
        "This appendix provides supplementary figures and tables that support the main manuscript. "
        "Terminology and measure labels are the same as in the main text.",
    )

    for item in items:
        add_caption(document, str(item["label"]), str(item["caption"]))
        if item["type"] == "figure":
            add_figure(document, Path(str(item["path"])))
        else:
            add_dataframe_table(document, item["data"])
            if item.get("merge_columns"):
                merge_context_cells(document.tables[-1], list(item["merge_columns"]))

    SUPPLEMENT_APPENDIX_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(SUPPLEMENT_APPENDIX_DOCX))


def build_markdown_appendix(items: list[dict[str, object]]) -> str:
    lines = [
        "# Supplementary Materials",
        "",
        "This appendix provides supplementary figures and tables that support the main manuscript. Terminology and measure labels are the same as in the main text.",
        "",
    ]
    for item in items:
        lines.extend([f"**{item['label']}** {item['caption']}", ""])
        if item["type"] == "figure":
            lines.extend([f"![{item['label']}]({Path(str(item['path'])).as_posix()})", ""])
        else:
            lines.extend([markdown_table(item["data"]), ""])
    return "\n".join(lines)


def build_inventory_items() -> list[dict[str, str]]:
    return [
        {
            "Item": "Figure S1",
            "Type": "Figure",
            "Title": "Calendar-year prevalence by application type",
            "Purpose": "Shows calendar-year patterning of the main, analytic, and explicit public RWE measures.",
            "Source file": str(SUPPLEMENT_FIGURE_DIR / "figure_s1_annual_trends.png"),
        },
        {
            "Item": "Figure S2",
            "Type": "Figure",
            "Title": "Evidence explicitness tier distribution",
            "Purpose": "Shows how directly the public record documented an observational real-world evidence basis for the labeling change.",
            "Source file": str(SUPPLEMENT_FIGURE_DIR / "figure_s2_explicitness_distribution.png"),
        },
        {
            "Item": "Table S1",
            "Type": "Table",
            "Title": "Calendar-year prevalence by application type",
            "Purpose": "Provides the counts underlying Figure S1.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s1_annual_trends.csv"),
        },
        {
            "Item": "Table S2",
            "Type": "Table",
            "Title": "Evidence explicitness tier counts",
            "Purpose": "Provides the counts underlying Figure S2.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s2_evidence_explicitness.csv"),
        },
        {
            "Item": "Table S3",
            "Type": "Table",
            "Title": "Sensitivity summary",
            "Purpose": "Summarizes prevalence under alternate measure definitions, subset restrictions, and provenance strata.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s3_sensitivity_summary.csv"),
        },
        {
            "Item": "Table S4",
            "Type": "Table",
            "Title": "Validation sample characteristics",
            "Purpose": "Describes the composition of the 300-event human validation sample.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s4_validation_sample_characteristics.csv"),
        },
        {
            "Item": "Table S5",
            "Type": "Table",
            "Title": "Product-age detail",
            "Purpose": "Shows descriptive prevalence by product age band and approval era.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s5_product_age_detail.csv"),
        },
        {
            "Item": "Table S6",
            "Type": "Table",
            "Title": "Therapeutic-area detail",
            "Purpose": "Shows descriptive heterogeneity across therapeutic areas.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s6_therapeutic_area_detail.csv"),
        },
        {
            "Item": "Table S7",
            "Type": "Table",
            "Title": "Sponsor/manufacturer detail",
            "Purpose": "Provides the exploratory sponsor/manufacturer comparisons.",
            "Source file": str(SUPPLEMENT_OUTPUT_DIR / "table_s7_sponsor_manufacturer_detail.csv"),
        },
    ]


def write_outputs(table_map: dict[str, pd.DataFrame], items: list[dict[str, object]]) -> None:
    SUPPLEMENT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    output_map = {
        "table_s1": "table_s1_annual_trends.csv",
        "table_s2": "table_s2_evidence_explicitness.csv",
        "table_s3": "table_s3_sensitivity_summary.csv",
        "table_s4": "table_s4_validation_sample_characteristics.csv",
        "table_s5": "table_s5_product_age_detail.csv",
        "table_s6": "table_s6_therapeutic_area_detail.csv",
        "table_s7": "table_s7_sponsor_manufacturer_detail.csv",
    }
    for key, filename in output_map.items():
        table_map[key].to_csv(SUPPLEMENT_OUTPUT_DIR / filename, index=False)

    inventory_items = build_inventory_items()
    inventory_df = build_inventory(inventory_items)
    inventory_df.to_csv(SUPPLEMENT_INVENTORY_CSV, index=False)
    SUPPLEMENT_INVENTORY_MD.write_text("# Supplement Inventory\n\n" + markdown_table(inventory_df) + "\n", encoding="utf-8")
    SUPPLEMENT_APPENDIX_MD.write_text(build_markdown_appendix(items), encoding="utf-8")
    build_docx_appendix(items)


def main() -> None:
    table_map = {
        "table_s1": clean_annual_trends(pd.read_csv(TABLE_PATHS["table_s1"])),
        "table_s2": clean_explicitness(pd.read_csv(TABLE_PATHS["table_s2"])),
        "table_s3": clean_sensitivity(pd.read_csv(TABLE_PATHS["table_s3"])),
        "table_s4": build_validation_sample_characteristics(pd.read_csv(VALIDATION_SAMPLE_PATH)),
        "table_s5": clean_product_age(pd.read_csv(TABLE_PATHS["table_s5"])),
        "table_s6": clean_therapeutic_area(pd.read_csv(TABLE_PATHS["table_s6"])),
        "table_s7": clean_sponsor_manufacturer(pd.read_csv(TABLE_PATHS["table_s7"])),
    }

    items = [
        {
            "type": "figure",
            "label": "Figure S1.",
            "caption": (
                "Calendar-year prevalence of broad public RWE, analytic public RWE, and explicit public RWE "
                "by application type. Note: percentages are calculated within each calendar-year and "
                "application-type stratum."
            ),
            "path": SUPPLEMENT_FIGURE_DIR / "figure_s1_annual_trends.png",
        },
        {
            "type": "figure",
            "label": "Figure S2.",
            "caption": (
                "Distribution of evidence explicitness tiers overall and by application type. Note: evidence explicitness "
                "tiers indicate how directly the public record documented an observational real-world evidence "
                "basis for the labeling change."
            ),
            "path": SUPPLEMENT_FIGURE_DIR / "figure_s2_explicitness_distribution.png",
        },
        {
            "type": "table",
            "label": "Table S1.",
            "caption": (
                "Calendar-year prevalence of broad public RWE, analytic public RWE, and explicit public RWE "
                "by application type. Note: this table provides the counts underlying Figure S1."
            ),
            "data": table_map["table_s1"],
            "merge_columns": [0, 1],
        },
        {
            "type": "table",
            "label": "Table S2.",
            "caption": (
                "Counts of evidence explicitness tiers overall and by application type. Note: this table "
                "provides the counts underlying Figure S2."
            ),
            "data": table_map["table_s2"],
            "merge_columns": [0],
        },
        {
            "type": "table",
            "label": "Table S3.",
            "caption": (
                "Sensitivity and provenance-stratified prevalence for the main and analytic public RWE measures. "
                "Note: positive values in the change column indicate higher prevalence than in the baseline full cohort."
            ),
            "data": table_map["table_s3"],
            "merge_columns": [0, 1],
        },
        {
            "type": "table",
            "label": "Table S4.",
            "caption": (
                "Characteristics of the 300-event human validation sample. Note: the validation sample was "
                "stratified for performance assessment and is not intended for direct prevalence estimation in the full cohort."
            ),
            "data": table_map["table_s4"],
            "merge_columns": [0],
        },
        {
            "type": "table",
            "label": "Table S5.",
            "caption": (
                "Product-age and approval-era distributions for broad public RWE and analytic public RWE. "
                "Note: percentages are event-level prevalences within each product-age or approval-era stratum."
            ),
            "data": table_map["table_s5"],
            "merge_columns": [0],
        },
        {
            "type": "table",
            "label": "Table S6.",
            "caption": (
                "Therapeutic-area distributions for broad public RWE and analytic public RWE. Note: therapeutic "
                "area was derived from FDA product metadata and used as a secondary heterogeneity layer."
            ),
            "data": table_map["table_s6"],
        },
        {
            "type": "table",
            "label": "Table S7.",
            "caption": (
                "Sponsor/manufacturer distributions for broad public RWE and analytic public RWE. Note: these "
                "groupings are heuristic metadata constructs and are presented as exploratory context rather than "
                "as a main explanatory analysis."
            ),
            "data": table_map["table_s7"],
            "merge_columns": [0],
        },
    ]

    write_outputs(table_map, items)


if __name__ == "__main__":
    main()
